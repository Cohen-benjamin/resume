"""Resume file -> plain text.

Handles the formats a resume actually arrives in. The repo's own resume is
HTML, which is the default, but a PDF export is the common case for everyone
else so all four are supported.
"""

from __future__ import annotations

import hashlib
import re
from pathlib import Path


class UnsupportedResume(ValueError):
    pass


def extract_text(path: Path) -> str:
    """Read a resume and return normalized plain text."""
    if not path.exists():
        raise FileNotFoundError(f"resume not found: {path}")

    suffix = path.suffix.lower()
    if suffix in {".html", ".htm"}:
        text = _from_html(path.read_text(encoding="utf-8", errors="replace"))
    elif suffix == ".pdf":
        text = _from_pdf(path)
    elif suffix == ".docx":
        text = _from_docx(path)
    elif suffix in {".txt", ".md", ""}:
        text = path.read_text(encoding="utf-8", errors="replace")
    else:
        raise UnsupportedResume(f"don't know how to read {suffix} resumes")

    return _tidy(text)


def _from_html(html: str) -> str:
    from selectolax.parser import HTMLParser

    tree = HTMLParser(html)
    for tag in tree.css("script, style, noscript"):
        tag.decompose()
    body = tree.body or tree.root
    # separator="\n" keeps each block element on its own line, which preserves
    # the bullet/heading structure the model uses to tell sections apart.
    return body.text(separator="\n") if body else ""


def _from_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _from_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _tidy(text: str) -> str:
    text = text.replace("\xa0", " ").replace("​", "")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    out: list[str] = []
    for line in lines:
        # Collapse runs of blank lines; HTML extraction produces a lot of them.
        if not line and out and not out[-1]:
            continue
        out.append(line)
    return "\n".join(out).strip()


def content_hash(text: str) -> str:
    return hashlib.sha256(text.encode()).hexdigest()[:16]
